#!/usr/bin/env python3
"""
更新報價表 Word 模板
- 簽署欄從兩欄改為三欄（加入批核人）
- 加入 {quoteMethodOther} 變數
- 檢查並修復字體大小
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

TEMPLATE_PATH = "public/templates/quotation.docx"
BACKUP_PATH = "public/templates/quotation_backup.docx"

def ensure_font_size(paragraph, min_size=6):
    """確保段落字體大小 >= min_size pt"""
    for run in paragraph.runs:
        if run.font.size and run.font.size < Pt(min_size):
            print(f"  ⚠️  修復過小字體：{run.font.size.pt}pt → 12pt")
            run.font.size = Pt(12)
        elif not run.font.size:
            run.font.size = Pt(12)

def find_signature_table(doc):
    """找到簽署欄表格（包含「索取報價人」）"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "索取報價人" in cell.text:
                    return table
    return None

def find_quote_method_section(doc):
    """找到報價方式區塊，返回該段落的索引"""
    for i, paragraph in enumerate(doc.paragraphs):
        if "報價方式" in paragraph.text or "quote_method" in paragraph.text.lower():
            return i
    return None

def update_signature_table(table):
    """將簽署表格從兩欄改為三欄"""
    print("\n🔧 更新簽署表格...")

    # 檢查現在有幾欄
    first_row = table.rows[0]
    num_cols = len(first_row.cells)
    print(f"  當前欄數：{num_cols}")

    if num_cols >= 3:
        print("  ✅ 已經是三欄，檢查內容...")
        # 檢查第三欄是否已有「批核人」
        if "批核人" in first_row.cells[2].text:
            print("  ✅ 批核人欄位已存在")
            return True

    # 如果只有兩欄，需要加入第三欄
    if num_cols == 2:
        print("  📝 加入第三欄（批核人）...")

        # 為每一行加入新欄位
        for row in table.rows:
            # 加入新 cell
            new_cell = row.add_cell()

        # 更新表頭
        header_row = table.rows[0]
        header_row.cells[0].text = "索取報價人"
        header_row.cells[1].text = "科組負責人"
        header_row.cells[2].text = "批核人"

        # 設定表頭樣式
        for cell in header_row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(12)

        # 更新簽署空白行（第二行）
        if len(table.rows) > 1:
            table.rows[1].cells[2].text = " "

        # 更新姓名職級行（第三行）
        if len(table.rows) > 2:
            name_row = table.rows[2]
            name_row.cells[2].text = "{approverName}  {approverRank}"

        # 更新日期行（第四行）
        if len(table.rows) > 3:
            date_row = table.rows[3]
            date_row.cells[2].text = "日期：{approverDate}"

        print("  ✅ 已加入批核人欄位")
        return True

    return False

def add_quote_method_other_field(doc, section_index):
    """在報價方式區塊後加入「其他報價方式說明」欄位"""
    print("\n🔧 加入「其他報價方式說明」欄位...")

    # 檢查是否已存在
    for para in doc.paragraphs:
        if "quoteMethodOther" in para.text:
            print("  ✅ 欄位已存在")
            return True

    if section_index is None:
        print("  ⚠️  找不到報價方式區塊，跳過")
        return False

    # 在報價方式段落後插入新段落
    # 注意：python-docx 不支援直接插入段落，需要在文件末尾加入
    # 這裡建議在模板中手動加入，或使用 docxtemplater 的條件顯示
    print("  ℹ️  建議在模板中手動加入 {quoteMethodOther} 變數")
    print("     位置：報價方式區塊後面")
    print("     格式：其他報價方式說明：{quoteMethodOther}")

    return False

def check_and_fix_fonts(doc):
    """檢查並修復所有段落的字體大小"""
    print("\n🔧 檢查字體大小...")
    fixed_count = 0

    # 檢查段落
    for para in doc.paragraphs:
        original_size = None
        for run in para.runs:
            if run.font.size and run.font.size < Pt(6):
                if original_size is None:
                    original_size = run.font.size.pt
                run.font.size = Pt(12)
                fixed_count += 1

        if original_size:
            print(f"  修復：{original_size}pt → 12pt")

    # 檢查表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.size and run.font.size < Pt(6):
                            run.font.size = Pt(12)
                            fixed_count += 1

    if fixed_count > 0:
        print(f"  ✅ 修復了 {fixed_count} 個過小字體")
    else:
        print("  ✅ 所有字體大小正常")

    return True

def main():
    print("=" * 60)
    print("更新報價表 Word 模板")
    print("=" * 60)

    # 檢查檔案是否存在
    if not os.path.exists(TEMPLATE_PATH):
        print(f"❌ 找不到模板檔案：{TEMPLATE_PATH}")
        return False

    print(f"\n📂 讀取模板：{TEMPLATE_PATH}")

    # 創建備份
    print(f"💾 創建備份：{BACKUP_PATH}")
    doc = Document(TEMPLATE_PATH)
    doc.save(BACKUP_PATH)

    # 重新載入
    doc = Document(TEMPLATE_PATH)

    # 1. 更新簽署表格
    sig_table = find_signature_table(doc)
    if sig_table:
        update_signature_table(sig_table)
    else:
        print("⚠️  找不到簽署表格")

    # 2. 加入其他報價方式欄位（提示）
    quote_section = find_quote_method_section(doc)
    add_quote_method_other_field(doc, quote_section)

    # 3. 檢查字體大小
    check_and_fix_fonts(doc)

    # 儲存更新後的模板
    print(f"\n💾 儲存更新後的模板：{TEMPLATE_PATH}")
    doc.save(TEMPLATE_PATH)

    print("\n" + "=" * 60)
    print("✅ 模板更新完成！")
    print("=" * 60)
    print(f"\n📋 備份檔案：{BACKUP_PATH}")
    print(f"📋 更新檔案：{TEMPLATE_PATH}")
    print("\n⚠️  請注意：")
    print("   1. {quoteMethodOther} 變數需要手動加入到報價方式區塊")
    print("   2. 建議在 Word 中打開檢查表格寬度是否需要調整")
    print("   3. 測試生成 DOCX 確認所有變數正確顯示")

    return True

if __name__ == "__main__":
    # 切換到專案根目錄
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    os.chdir(project_root)

    success = main()
    exit(0 if success else 1)
