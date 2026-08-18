#!/usr/bin/env python3
"""
驗證模板更新結果
"""
from docx import Document

doc = Document("public/templates/quotation.docx")

print("=" * 60)
print("模板更新驗證")
print("=" * 60)

# 找到簽署欄段落
print("\n📝 簽署欄段落（索引 9-20）：")
for i in range(9, min(21, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if text:
        print(f"\n[{i}] {text[:100]}")
        if i == 10:
            # 檢查是否包含所有必要變數
            required_vars = [
                "requestorName", "requestorRank", "requestorDate",
                "deptHeadName", "deptHeadRank", "deptHeadDate",
                "approverName", "approverRank", "approverDate"
            ]
            found = [v for v in required_vars if v in text]
            print(f"    ✓ 包含變數：{', '.join(found)}")
            missing = [v for v in required_vars if v not in text]
            if missing:
                print(f"    ✗ 缺少變數：{', '.join(missing)}")
            else:
                print(f"    ✅ 所有簽署變數都已存在")

print("\n" + "=" * 60)
print("✅ 驗證完成")
print("=" * 60)
