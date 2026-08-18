#!/usr/bin/env python3
"""
更新報價表 Word 模板 - 處理段落格式的簽署欄
"""
import os
from docx import Document
from docx.shared import Pt
from copy import deepcopy

TEMPLATE_PATH = "public/templates/quotation.docx"
BACKUP_PATH = "public/templates/quotation_backup.docx"

def find_signature_paragraphs(doc):
    """找到簽署欄的段落索引"""
    indices = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "索取報價人" in text and "姓名" in text:
            indices['header'] = i
        elif "{requestorName}" in text:
            indices['requestor'] = i
        elif "{deptHeadName}" in text:
            indices['deptHead'] = i
    return indices

def update_signature_section(doc, indices):
    """更新簽署段落，加入批核人"""
    print("\n🔧 更新簽署欄...")

    if 'header' not in indices:
        print("  ⚠️  找不到簽署欄表頭")
        return False

    # 取得現有段落
    header_para = doc.paragraphs[indices['header']]
    print(f"  📝 表頭內容：{header_para.text[:60]}")

    # 檢查是否已經有批核人
    if "批核人" in header_para.text:
        print("  ✅ 已包含批核人欄位")
        # 檢查資料行
        if 'deptHead' in indices:
            dept_para = doc.paragraphs[indices['deptHead']]
            if "{approverName}" in dept_para.text:
                print("  ✅ 批核人資料行已存在")
                return True

    # 更新表頭：索取報價人 → 索取報價人、科組負責人、批核人
    # 原格式：索取報價人 姓名   	職級     	簽署	日期
    # 改為三組
    header_para.text = "索取報價人 姓名   \t職級     \t簽署\t日期\t\t科組負責人 姓名   \t職級     \t簽署\t日期\t\t批核人 姓名   \t職級     \t簽署\t日期"

    # 設定字體
    for run in header_para.runs:
        run.font.bold = True
        run.font.size = Pt(12)

    print("  ✅ 已更新表頭")

    # 更新資料行
    if 'requestor' in indices:
        requestor_para = doc.paragraphs[indices['requestor']]
        # 原格式：{requestorName}\t{requestorRank}\t\t{requestorDate}
        # 改為三組
        requestor_para.text = "{requestorName}\t{requestorRank}\t\t{requestorDate}\t\t{deptHeadName}\t{deptHeadRank}\t\t{deptHeadDate}\t\t{approverName}\t{approverRank}\t\t{approverDate}"

        for run in requestor_para.runs:
            run.font.size = Pt(12)

        print("  ✅ 已更新資料行")

    # 刪除舊的科組負責人段落（如果存在）
    if 'deptHead' in indices:
        # 找到並標記為刪除（python-docx 不支援直接刪除段落）
        # 改為清空內容
        dept_para = doc.paragraphs[indices['deptHead']]
        # 檢查是否有獨立的科組負責人段落
        if indices['deptHead'] == indices['requestor'] + 2:
            # 有獨立的科組負責人表頭和資料行，需要刪除
            print("  ℹ️  偵測到舊的兩欄格式，建議手動刪除多餘段落")

    return True

def add_quote_method_other(doc):
    """在報價方式後加入其他說明欄位"""
    print("\n🔧 加入「其他報價方式說明」欄位...")

    # 找到報價方式段落
    quote_method_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower()
        if "報價方式" in para.text or "quotemethod" in text:
            quote_method_idx = i
            break

    if quote_method_idx is None:
        print("  ⚠️  找不到報價方式段落")
        return False

    # 檢查是否已有 quoteMethodOther
    for para in doc.paragraphs:
        if "quoteMethodOther" in para.text:
            print("  ✅ 欄位已存在")
            return True

    # 找到插入位置（報價方式後的幾個段落）
    insert_after = quote_method_idx + 3  # 跳過報價方式的 checkbox 行

    # 創建新段落（在指定位置後插入）
    # python-docx 限制：只能在文件末尾加入，無法指定位置插入
    # 建議手動在模板中加入
    print("  ℹ️  建議手動在報價方式區塊後加入：")
    print("     其他報價方式說明：{quoteMethodOther}")

    return False

def check_and_fix_fonts(doc):
    """檢查並修復所有段落的字體大小"""
    print("\n🔧 檢查字體大小...")
    fixed_count = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size and run.font.size < Pt(6):
                run.font.size = Pt(12)
                fixed_count += 1
            elif not run.font.size:
                run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.size and run.font.size < Pt(6):
                            run.font.size = Pt(12)
                            fixed_count += 1

    if fixed_count > 0:
        print(f"  ✅ 修復了 {fixed_count} 個過小字體")
    else:
        print("  ✅ 所有字體大小正常")

    return True

def main():
    print("=" * 60)
    print("更新報價表 Word 模板（段落格式）")
    print("=" * 60)

    if not os.path.exists(TEMPLATE_PATH):
        print(f"❌ 找不到模板檔案：{TEMPLATE_PATH}")
        return False

    print(f"\n📂 讀取模板：{TEMPLATE_PATH}")

    # 創建備份（如果還沒有）
    if not os.path.exists(BACKUP_PATH):
        print(f"💾 創建備份：{BACKUP_PATH}")
        doc = Document(TEMPLATE_PATH)
        doc.save(BACKUP_PATH)
    else:
        print(f"💾 備份已存在：{BACKUP_PATH}")

    # 載入模板
    doc = Document(TEMPLATE_PATH)

    # 1. 找到並更新簽署欄
    indices = find_signature_paragraphs(doc)
    print(f"\n🔍 找到的段落索引：{indices}")
    update_signature_section(doc, indices)

    # 2. 加入其他報價方式欄位
    add_quote_method_other(doc)

    # 3. 檢查字體
    check_and_fix_fonts(doc)

    # 儲存
    print(f"\n💾 儲存更新後的模板：{TEMPLATE_PATH}")
    doc.save(TEMPLATE_PATH)

    print("\n" + "=" * 60)
    print("✅ 模板更新完成！")
    print("=" * 60)
    print(f"\n📋 備份檔案：{BACKUP_PATH}")
    print(f"📋 更新檔案：{TEMPLATE_PATH}")
    print("\n⚠️  請在 Word 中檢查：")
    print("   1. 簽署欄三欄對齊是否正確（可能需要調整 Tab 位置）")
    print("   2. 手動加入 {quoteMethodOther} 變數到報價方式區塊")
    print("   3. 刪除多餘的科組負責人段落（如果有）")

    return True

if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    os.chdir(project_root)

    success = main()
    exit(0 if success else 1)
