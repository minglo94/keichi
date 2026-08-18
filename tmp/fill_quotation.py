"""
fill_quotation.py
Fills 報價表Oral Quotation2425.docx template with provided data.
Run structure mapped from template inspection (DT version, 4-row x 7-col table).
"""

import glob
import os
from copy import deepcopy
from datetime import datetime

from docx import Document

# ── Constants ─────────────────────────────────────────────────────────────────
CHECKED = '\uf052'    # Wingdings 2 checked checkbox (checkmark in box)
UNCHECKED = '\uf0a3'  # Wingdings 2 empty checkbox (empty box)

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'output')


def _find_template():
    """
    Locate the quotation template.
    Priority: (1) template_base.docx bundled beside this script,
              (2) glob search on OneDrive (original machine only).
    """
    # 1. Bundled local copy — works on any machine after distribution
    local = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'template_base.docx')
    if os.path.isfile(local):
        return local

    # 2. Fall back to OneDrive glob (original development machine)
    candidates = glob.glob(
        'C:/Users/CMLO/OneDrive - ccckcss/**/報價表Oral Quotation2425.docx',
        recursive=True,
    )
    if not candidates:
        raise FileNotFoundError(
            'Template not found. Place template_base.docx beside app.py.'
        )
    no_amp = [c for c in candidates if '&' not in c and 'new' in c]
    if no_amp:
        return no_amp[0]
    no_amp2 = [c for c in candidates if '&' not in c]
    return no_amp2[0] if no_amp2 else candidates[0]


def _find_fill_para(doc, label_text):
    """
    Return the paragraph immediately following the one that contains label_text.
    More robust than hard-coded indices when template minor versions differ.
    """
    for i, p in enumerate(doc.paragraphs[:-1]):
        if label_text in p.text:
            return doc.paragraphs[i + 1]
    return None


def _clear_run(run, text=''):
    """Set run text, preserve font metadata."""
    run.text = text


def _fill_cell_para(para, text):
    """Replace all runs in a paragraph with a single text value."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for r in para.runs[1:]:
        r.text = ''


def _set_cell_lines(cell, lines):
    """
    Set cell content to a list of lines (one per paragraph).
    Uses existing paragraphs for lines 0-N, leaves extra paras blank.
    """
    paras = cell.paragraphs
    for i, line in enumerate(lines):
        if i < len(paras):
            _fill_cell_para(paras[i], str(line))
        else:
            # Add new paragraph copying format from first
            new_elem = deepcopy(paras[0]._element)
            cell._element.append(new_elem)
            _fill_cell_para(cell.paragraphs[-1], str(line))
    # Clear remaining paragraphs
    for j in range(len(lines), len(paras)):
        _fill_cell_para(paras[j], '')


def fill_shared_cells(table, items, quotation_name):
    """
    Fill the merged item-description and quantity cells (columns 1 & 2).
    These cells are vertically merged across all supplier rows — must be filled
    only ONCE using rows[1] (the top visible cell).
    """
    ref_row = table.rows[1]

    # Cell[1]: quotation name (para 0) + item names (paras 1-3)
    cell1 = ref_row.cells[1]
    _fill_cell_para(cell1.paragraphs[0], quotation_name)
    for idx in range(3):
        item_name = items[idx]['name'] if idx < len(items) else ''
        if idx + 1 < len(cell1.paragraphs):
            _fill_cell_para(cell1.paragraphs[idx + 1], item_name)

    # Cell[2]: quantities (para 0 is blank header-align, paras 1-3 are quantities)
    cell2 = ref_row.cells[2]
    if cell2.paragraphs:
        _fill_cell_para(cell2.paragraphs[0], '')
    for idx in range(3):
        qty = str(items[idx]['qty']) if idx < len(items) else ''
        if idx + 1 < len(cell2.paragraphs):
            _fill_cell_para(cell2.paragraphs[idx + 1], qty)


def fill_supplier_row(row, supplier_data, is_recommended):
    """Fill supplier-specific cells (columns 3-6) for one supplier row.
    Does NOT touch columns 1-2 (vertically merged shared cells).
    """
    items = supplier_data.get('_items', [])

    # Cell[3]: supplier name (para 0) + tel (para 1)
    cell3 = row.cells[3]
    if cell3.paragraphs:
        _fill_cell_para(cell3.paragraphs[0], supplier_data.get('name', ''))
    if len(cell3.paragraphs) > 1:
        _fill_cell_para(cell3.paragraphs[1], supplier_data.get('tel', ''))

    # Cell[4]: unit prices — one price per item, separate paragraphs
    cell4 = row.cells[4]
    prices = supplier_data.get('prices', [])
    n_items = max(len(items), 1)
    price_lines = [str(prices[i]) if i < len(prices) else '' for i in range(n_items)]
    _set_cell_lines(cell4, price_lines)

    # Cell[5]: total price
    cell5 = row.cells[5]
    _fill_cell_para(cell5.paragraphs[0], str(supplier_data.get('total', '')))

    # Cell[6]: adoption checkmark
    cell6 = row.cells[6]
    mark = '✓' if is_recommended else ''
    if cell6.paragraphs:
        p = cell6.paragraphs[0]
        if p.runs:
            p.runs[0].text = mark
            for r in p.runs[1:]:
                r.text = ''
        else:
            p.add_run(mark)


def fill_quotation(data: dict) -> str:
    """
    Fill the quotation template with the provided data dict.

    Required keys in data:
        quotation_date      str  e.g. '2026-05-03'
        quote_method        str  'phone'|'fax'|'mail'|'other'
        quotation_name      str  e.g. '無線鍵盤及滑鼠'
        items               list of {'name': str, 'qty': int/str}
        supplier_a          dict {'name', 'tel', 'prices': list, 'total'}
        supplier_b          dict {'name', 'tel', 'prices': list, 'total'}
        recommended         str  'A'|'B'
        use_lower_price     bool
        higher_price_reason str  (optional)
        fewer_suppliers_reason str (optional)
        item_category       str  'fixed'|'consumable'|'other'
        category_other      str  (if item_category == 'other')
        department          str
        purpose             str
        delivery_date       str
        funding_source      str
        requestor_name      str
        requestor_rank      str
        requestor_date      str
        dept_head_name      str
        dept_head_rank      str
        dept_head_date      str

    Returns path to generated DOCX file.
    """
    src = _find_template()
    doc = Document(src)

    recommended = data.get('recommended', 'A')
    items = data.get('items', [{'name': '', 'qty': ''}])
    supplier_a = data.get('supplier_a', {})
    supplier_b = data.get('supplier_b', {})

    # ── Para[3]: date / method / recommended supplier ─────────────────────────
    p3 = doc.paragraphs[3]
    runs = p3.runs

    # Date (run[1])
    q_date = data.get('quotation_date', '')
    if q_date:
        try:
            # Parse YYYY-MM-DD or similar and convert to traditional Chinese date format
            for fmt in ('%Y-%m-%d', '%Y/%m/%d'):
                try:
                    dt = datetime.strptime(q_date, fmt)
                    date_str = f"  {dt.year} 年    {dt.month} 月  {dt.day}   日"
                    break
                except ValueError:
                    continue
            else:
                date_str = q_date
        except Exception:
            date_str = q_date
    else:
        date_str = "___________"
    if len(runs) > 1:
        _clear_run(runs[1], date_str)

    # Quotation method checkboxes (runs 5/8/11/14)
    method = data.get('quote_method', 'phone')
    method_map = {5: 'phone', 8: 'fax', 11: 'mail', 14: 'other'}
    for ridx, key in method_map.items():
        if ridx < len(runs):
            runs[ridx].text = CHECKED if method == key else UNCHECKED

    # Other method description (run[18])
    if len(runs) > 18:
        if method == 'other':
            _clear_run(runs[18], f"  {data.get('quote_method_other', '')}  ")
        else:
            _clear_run(runs[18], "    ")

    # Recommended supplier name (run[33])
    rec_name = supplier_a.get('name', '') if recommended == 'A' else supplier_b.get('name', '')
    if len(runs) > 33:
        _clear_run(runs[33], f"        {rec_name}    ")

    # Price type checkboxes (run[39]=lower, run[43]=higher)
    use_lower = data.get('use_lower_price', True)
    if len(runs) > 39:
        runs[39].text = CHECKED if use_lower else UNCHECKED
    if len(runs) > 43:
        runs[43].text = CHECKED if not use_lower else UNCHECKED

    # Higher price reason (run[47])
    if len(runs) > 47:
        _clear_run(runs[47], data.get('higher_price_reason', ''))

    # ── Para[5]: category / dept / purpose / delivery / funding ──────────────
    p5 = doc.paragraphs[5]
    r5 = p5.runs

    # Fewer suppliers reason (runs 5-6)
    if len(r5) > 6:
        _clear_run(r5[5], data.get('fewer_suppliers_reason', ''))
        _clear_run(r5[6], '')

    # Item category checkboxes (runs 17/20/23)
    cat = data.get('item_category', 'consumable')
    cat_map = {17: 'fixed', 20: 'consumable', 23: 'other'}
    for ridx, key in cat_map.items():
        if ridx < len(r5):
            r5[ridx].text = CHECKED if cat == key else UNCHECKED

    # Category other description (runs 26-28)
    cat_other = data.get('category_other', '') if cat == 'other' else ''
    if len(r5) > 28:
        _clear_run(r5[26], cat_other)
        _clear_run(r5[27], '')
        _clear_run(r5[28], '')

    # Department (runs 31-32)
    dept = data.get('department', '')
    if len(r5) > 32:
        _clear_run(r5[31], dept)
        _clear_run(r5[32], '')

    # Purpose (runs 36-38)
    purpose = data.get('purpose', '')
    if len(r5) > 38:
        _clear_run(r5[36], purpose)
        _clear_run(r5[37], '')
        _clear_run(r5[38], '')

    # Delivery date (run[51])
    if len(r5) > 51:
        _clear_run(r5[51], data.get('delivery_date', ''))

    # Funding source (runs 58-59)
    if len(r5) > 59:
        _clear_run(r5[58], data.get('funding_source', ''))
        _clear_run(r5[59], '')

    # ── Requestor & dept-head fill-in lines ───────────────────────────────────
    today = datetime.now().strftime('%Y-%m-%d')

    # Locate paragraphs by content (robust across template minor versions)
    def _find_label_para(doc, label_text):
        for p in doc.paragraphs:
            if label_text in p.text:
                return p
        return None

    lbl_p9 = _find_label_para(doc, '索取報價人') or doc.paragraphs[9]
    lbl_p16 = _find_label_para(doc, '科組負責人') or doc.paragraphs[16]
    lbl_p26 = _find_label_para(doc, '批核人') or doc.paragraphs[26]

    val_p10 = _find_fill_para(doc, '索取報價人') or doc.paragraphs[10]
    val_p17 = _find_fill_para(doc, '科組負責人') or doc.paragraphs[17]
    val_p27 = _find_fill_para(doc, '批核人') or doc.paragraphs[27]

    def _fill_sig_row(para, label_prefix, name, rank, date, is_approver=False):
        # Save formatting from the first run if available to maintain font styling
        font_name = None
        font_size = None
        bold = None
        italic = None
        if para.runs:
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.font.bold
            italic = first_run.font.italic
            
        # Clear all runs
        p_element = para._element
        for r in list(para.runs):
            p_element.remove(r._element)
        para.runs.clear()
        
        # Helper to add run with copied formatting
        from docx.shared import Pt
        def add_formatted_run(text, is_value=False):
            run = para.add_run(text)
            if font_name:
                run.font.name = font_name
            # If font size is missing or extremely small, default to 12pt
            if font_size and font_size >= Pt(6):
                run.font.size = font_size
            else:
                run.font.size = Pt(12)
            # Make label bold and value non-bold for clear contrast
            if is_value:
                run.font.bold = False
            else:
                run.font.bold = bold if bold is not None else True
            if italic is not None:
                run.font.italic = italic
            return run

        # Reconstruct the paragraph in one row
        add_formatted_run(f"{label_prefix} 姓名: ")
        add_formatted_run(name, is_value=True)
        add_formatted_run("\t職級: ")
        add_formatted_run(rank, is_value=True)
        if is_approver:
            add_formatted_run("\t校長/副校長")
        add_formatted_run("\t簽署: ")
        add_formatted_run("\t日期: ")
        add_formatted_run(date, is_value=True)

    def _clear_paragraph(para):
        p_element = para._element
        for r in list(para.runs):
            p_element.remove(r._element)
        para.runs.clear()
        # Set line height and space before/after to 0/1pt to shrink paragraph completely
        from docx.shared import Pt
        para.paragraph_format.line_spacing = Pt(1)
        para.paragraph_format.line_spacing_rule = None
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    _fill_sig_row(
        lbl_p9,
        label_prefix="索取報價人",
        name=data.get('requestor_name', ''),
        rank=data.get('requestor_rank', ''),
        date=data.get('requestor_date') or today,
    )
    _clear_paragraph(val_p10)

    _fill_sig_row(
        lbl_p16,
        label_prefix="科組負責人",
        name=data.get('dept_head_name', ''),
        rank=data.get('dept_head_rank', ''),
        date=data.get('dept_head_date') or today,
    )
    _clear_paragraph(val_p17)

    _fill_sig_row(
        lbl_p26,
        label_prefix="批核人",
        name=data.get('approver_name', ''),
        rank=data.get('approver_rank', ''),
        date=data.get('approver_date') or today,
        is_approver=True,
    )
    _clear_paragraph(val_p27)


    # ── Table rows ─────────────────────────────────────────────────────────────
    table = doc.tables[0]
    quotation_name = data.get('quotation_name', '')

    # Fill the vertically-merged item/quantity cells once via the first supplier row
    fill_shared_cells(table, items, quotation_name)

    # Attach items to supplier dicts so fill_supplier_row can compute price count
    supplier_a['_items'] = items
    supplier_b['_items'] = items

    # Fill supplier-specific columns (3-6) for rows 1 & 2
    fill_supplier_row(table.rows[1], supplier_a, recommended == 'A')
    fill_supplier_row(table.rows[2], supplier_b, recommended == 'B')

    # Row 3 (optional 3rd supplier): clear only supplier-specific cells
    if len(table.rows) > 3:
        fill_supplier_row(table.rows[3], {'name': '', 'tel': '', 'prices': [], 'total': '', '_items': items}, False)

    # ── Save output ────────────────────────────────────────────────────────────
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    dept_safe = data.get('department', '科組').replace('/', '_').replace('\\', '_').replace(' ', '_')
    out_name = f'報價表_{ts}_{dept_safe}.docx'
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    return out_path
