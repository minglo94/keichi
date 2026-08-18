#!/usr/bin/env python3
"""
清理模板中重複的科組負責人段落
"""
from docx import Document
from docx.oxml import OxmlElement

def delete_paragraph(paragraph):
    """刪除段落"""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def clean_duplicate_sections(doc):
    """清理重複的科組負責人段落"""
    print("\n🔧 清理重複段落...")

    # 找到需要刪除的段落（索引 16-17 附近的舊科組負責人）
    to_delete = []

    for i in range(11, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # 如果是獨立的科組負責人段落（不是合併後的）
        if i > 10 and "科組負責人 姓名" in text and "索取報價人" not in text:
            print(f"  ⚠️  找到重複的科組負責人表頭（段落 {i}）")
            to_delete.append(i)
            # 也標記下一行（資料行）
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if "{deptHeadName}" in next_para.text and "{requestorName}" not in next_para.text:
                    to_delete.append(i + 1)
                    print(f"  ⚠️  找到重複的科組負責人資料行（段落 {i+1}）")

    if not to_delete:
        print("  ✅ 沒有重複段落")
        return True

    # 從後往前刪除（避免索引變化）
    for idx in sorted(to_delete, reverse=True):
        para = doc.paragraphs[idx]
        print(f"  🗑️  刪除段落 {idx}：{para.text.strip()[:50]}")
        delete_paragraph(para)

    print(f"  ✅ 已刪除 {len(to_delete)} 個重複段落")
    return True

def main():
    print("=" * 60)
    print("清理模板重複段落")
    print("=" * 60)

    doc = Document("public/templates/quotation.docx")

    clean_duplicate_sections(doc)

    print(f"\n💾 儲存清理後的模板")
    doc.save("public/templates/quotation.docx")

    print("\n✅ 清理完成！")

if __name__ == "__main__":
    import os
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    os.chdir(project_root)

    main()
