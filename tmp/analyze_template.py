#!/usr/bin/env python3
"""
分析報價表模板結構
"""
from docx import Document

doc = Document("public/templates/quotation.docx")

print("=" * 60)
print("模板結構分析")
print("=" * 60)

print(f"\n📊 統計：")
print(f"  段落數：{len(doc.paragraphs)}")
print(f"  表格數：{len(doc.tables)}")

print(f"\n📝 段落內容預覽（最後 20 個）：")
for i, para in enumerate(doc.paragraphs[-20:], start=len(doc.paragraphs)-20):
    text = para.text.strip()
    if text:
        print(f"  [{i}] {text[:80]}")

print(f"\n📊 表格結構：")
for i, table in enumerate(doc.tables):
    print(f"\n表格 {i+1}:")
    print(f"  行數：{len(table.rows)}")
    print(f"  欄數：{len(table.rows[0].cells) if table.rows else 0}")

    # 預覽前幾行
    for row_idx, row in enumerate(table.rows[:3]):
        cells_text = [cell.text.strip()[:30] for cell in row.cells]
        print(f"    行 {row_idx+1}: {' | '.join(cells_text)}")

print("\n🔍 搜尋關鍵字「索取報價人」：")
found = False
for i, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if "索取報價人" in cell.text or "requestor" in cell.text.lower():
                print(f"  ✓ 找到於表格 {i+1}，行 {row_idx+1}，欄 {cell_idx+1}")
                print(f"    內容：{cell.text.strip()}")
                found = True

if not found:
    print("  ⚠️  未在表格中找到")
    print("\n🔍 在段落中搜尋：")
    for i, para in enumerate(doc.paragraphs):
        if "索取報價人" in para.text or "requestor" in para.text.lower():
            print(f"  ✓ 找到於段落 {i}")
            print(f"    內容：{para.text.strip()}")
